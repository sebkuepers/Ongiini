#!/usr/bin/env python3
"""Generate a phone-friendly Word document for Elizabeth (the hired
Oshiwambo translator) to fill in ground-truth translations.

Reads data/oshiwambo_eval_v2.tsv (~400 items, 5 domains, phenomenon-
tagged). Produces a .docx with EN source only — Claude / Gemma machine
translations are deliberately hidden so Elizabeth's work isn't anchored
on them. Items are randomised so domain ordering doesn't bias her.

Layout decisions for mobile Word editing:
- No tables (mobile Word renders them painfully).
- Each phrase is its own block with predictable structure so her thumb
  knows where to land: bold label, then a blank line.
- The phrase ID is shown so she can refer back to specific items.
- The blind-split flag is NOT shown to her — that's our reporting
  bookkeeping, not her concern.
"""
from __future__ import annotations

import argparse
import csv
import random
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor

DEFAULT_SOURCE = (
    Path(__file__).resolve().parents[1] / "data/oshiwambo_eval_v2.tsv"
)
DEFAULT_OUT = (
    Path.home() / "Desktop/ongiini-eval-elizabeth-v2.docx"
)


def load_rows(path: Path) -> list[dict]:
    rows: list[dict] = []
    with path.open() as f:
        for r in csv.DictReader(f, delimiter="\t"):
            en = (r.get("english") or "").strip()
            if not en:
                continue
            rows.append(r)
    return rows


def add_label(p, text: str, bold: bool = True, size: int = 12) -> None:
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def build_doc(rows: list[dict], shuffle_seed: int) -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)

    # ── Title + intro ───────────────────────────────────────────
    doc.add_heading(
        "Ongiini AI — Oshindonga & Oshikwanyama eval (v2)", level=0
    )

    intro = doc.add_paragraph()
    intro.add_run(
        f"Tangi unene! Thank you for helping us. Below are "
        f"{len(rows)} short English sentences. For each one, please "
        "write the natural, everyday way a Namibian would say it in "
        "both Oshindonga AND Oshikwanyama."
    )

    intro2 = doc.add_paragraph()
    intro2.add_run(
        "You can edit this document on your phone. Tap the blank line "
        "below \"Oshindonga:\" or \"Oshikwanyama:\" and start typing. "
        "Save as you go — no need to finish in one sitting."
    )

    intro3 = doc.add_paragraph()
    intro3.add_run(
        "If a sentence doesn't translate naturally in one of the "
        "languages, just write \"n/a\" and a short note about why "
        "(e.g. \"only used in Oshindonga\", \"no direct equivalent\")."
    )

    intro4 = doc.add_paragraph()
    note_run = intro4.add_run(
        "Please write naturally — the way you would actually speak or "
        "text someone. Pick the most common, neutral register (not the "
        "most formal or textbook version, but also not the most casual "
        "slang). When two valid translations exist, pick the one a "
        "wider audience would understand."
    )
    note_run.italic = True

    intro5 = doc.add_paragraph()
    intro5.add_run(
        "Don't worry about getting things \"right\" — your job is to "
        "tell us how Namibians actually express each idea. Your "
        "translation IS the standard; we're measuring our AI against "
        "you, not the other way around."
    )

    # ── Phrases (shuffled so domain ordering doesn't bias) ───────
    rng = random.Random(shuffle_seed)
    ordered = list(rows)
    rng.shuffle(ordered)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("The 425 phrases", level=1)

    total = len(ordered)
    for counter, r in enumerate(ordered, start=1):
        # Tiny grey counter at top of each block
        counter_p = doc.add_paragraph()
        cr = counter_p.add_run(f"Phrase {counter} of {total}  ·  id #{r['id']}")
        cr.bold = True
        cr.font.size = Pt(10)
        cr.font.color.rgb = RGBColor(0x6c, 0x6c, 0x6c)

        # English text
        en_p = doc.add_paragraph()
        en_label = en_p.add_run("English:  ")
        en_label.bold = True
        en_p.add_run(r["english"])

        # Oshindonga label + blank line
        odg_p = doc.add_paragraph()
        add_label(odg_p, "Oshindonga:")
        doc.add_paragraph()

        # Oshikwanyama label + blank line
        okw_p = doc.add_paragraph()
        add_label(okw_p, "Oshikwanyama:")
        doc.add_paragraph()

        # Visual separator
        sep = doc.add_paragraph("─" * 25)
        sep.runs[0].font.color.rgb = RGBColor(0xc0, 0xc0, 0xc0)

    # ── Closing ──────────────────────────────────────────────────
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("That's it — tangi unene!", level=1)
    doc.add_paragraph(
        "When you're done (or whenever you want to share what you "
        "have so far), please save the file and send it back. We'll "
        "review and let you know if anything needs another look. "
        "If a phrase had something strange about it or you weren't "
        "sure about it, you can also write a short note in the file."
    )

    return doc


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--source", default=str(DEFAULT_SOURCE),
                    help="Path to v2 TSV. Default: data/oshiwambo_eval_v2.tsv")
    ap.add_argument("--out", default=str(DEFAULT_OUT),
                    help="Output .docx path. Default: ~/Desktop/...v2.docx")
    ap.add_argument("--seed", type=int, default=42,
                    help="Random seed for phrase shuffle (so re-runs are stable)")
    args = ap.parse_args(argv)

    source = Path(args.source)
    out = Path(args.out)
    if not source.exists():
        print(f"ERROR: source not found at {source}", file=sys.stderr)
        return 1
    rows = load_rows(source)
    print(f"loaded {len(rows)} phrases from {source.name}", file=sys.stderr)
    doc = build_doc(rows, args.seed)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"wrote {out}  ({out.stat().st_size / 1024:.1f} KB)", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
